from docx import Document
from docx.shared import Inches

def generate_word_document(image_path, text, output_filename):
    """
    Generates a DOCX file with both text and an image.

    Args:
    image_path (str): The path to the image file.
    text (str): The text to be included in the document.
    output_filename (str): The name of the output DOCX file.
    """
    # Create a new Document object
    doc = Document()

    # Add text to the document
    doc.add_paragraph(text)
    
    # Add an image with proper sizing
    try:
        doc.add_picture(image_path, width=Inches(4.0))  # Adjust image size as needed
    except Exception as e:
        print(f"Error adding image: {e}")
    
    # Save the document
    try:
        doc.save(output_filename)
        print(f"Document saved as {output_filename}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    # Example usage: Modify this part as necessary for your needs
    generate_word_document('path_to_your_image.jpg', 'This is an example text.', 'output.docx')



# from docx import Document
# from docx.shared import Inches
# from PIL import Image
# import os

# def generate_word_document(text, image_path, output_path):
#     try:
#         # Create a Word document
#         doc = Document()

#         # Add a title
#         doc.add_heading("Draft Blog Post", level=1)

#         # Add the text
#         doc.add_heading("Extracted Text", level=2)
#         doc.add_paragraph(text)

#         # Add the image if provided
#         if image_path and os.path.exists(image_path):
#             doc.add_heading("Associated Image", level=2)

#             # Resize the image to fit the Word document
#             img = Image.open(image_path)
#             img_width, img_height = img.size
#             max_width = 6  # Max width in inches
#             scaling_factor = min(1, max_width / img_width)
#             resized_width = int(img_width * scaling_factor)
#             resized_height = int(img_height * scaling_factor)
#             img = img.resize((resized_width, resized_height))

#             # Save resized image to a temporary file
#             temp_img_path = os.path.join(os.path.dirname(output_path), "resized_image.png")
#             img.save(temp_img_path)

#             # Add the resized image to the document
#             doc.add_picture(temp_img_path, width=Inches(resized_width / 96))  # 96 DPI conversion

#         # Save the document
#         doc.save(output_path)
#         return True

#     except Exception as e:
#         print(f"Error generating Word document: {e}")
#         return False

